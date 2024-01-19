from docx import Document
import os
import re
from datetime import time, datetime, timedelta
from pprint import pprint


def replace_text_in_doc(doc, replacer):
    def replace_text_in_paragraph(paragraph):
        if not paragraph.runs:
            return paragraph

        text = paragraph.text
        if '$' not in text and '(' not in text:
            return

        # Combine other runs into the first run
        first_run = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            first_run.text += run.text
            p = paragraph._p
            p.remove(run._r)

        replacer(first_run)

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph)

def add_minutes(strtime, minutes):
    tm = time.fromisoformat(strtime)
    fulldate = datetime(100, 1, 1, tm.hour, tm.minute, tm.second)
    fulldate = fulldate + timedelta(minutes=minutes)
    return fulldate.strftime('%H:%M')

def handle_no(no):
    if no.endswith('1'):
        no += 'st'
    elif no.endswith('2'):
        no += 'nd'
    elif no.endswith('3'):
        no += 'rd'
    else:
        no += 'th'
    return no

data_table = {}

def handle_agenda(doc):
    last_session = add_minutes(data_table['TB'], 10)
    session_minutes = 0

    def replacer(item):
        print('<<<', item.text)
        def repl(match):
            label = match.group(1)
            if label in data_table:
                if label == 'NO':
                    return handle_no(data_table['NO'])
                return data_table[label]

            nonlocal last_session
            nonlocal session_minutes

            minutes = match.group(3)
            if minutes:                
                session_minutes += int(minutes)
            print('minutes', minutes, session_minutes)

            if label == 'TS':
                last_session = add_minutes(last_session, session_minutes)
                session_minutes = 0
                return last_session
            elif label == 'TSS':
                last_session = add_minutes(last_session, session_minutes+1)
                session_minutes = 0
                return last_session

            return match.group(0)

        item.text = re.sub(r'\$\{(\w+)\}|\((\d+\'\-)?(\d+)\'\)', repl, item.text)
        print(">>>", item.text)

    replace_text_in_doc(doc, replacer)

    data_table['TE'] = add_minutes(last_session, session_minutes)
    replace_text_in_doc(doc, replacer)

def handle_meeting_info(doc):
    for table in doc.tables:
        row = table.rows[0]
        cell = row.cells[0]
        title = cell.text.strip()
        print(title)
        if title != 'Meeting Information':
            continue
        for row in table.rows[1:]:
            label = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            if not label:
                continue
            data_table[label] = value
            print(label, ":", value)
        # remove table
        table._element.getparent().remove(table._element)
    print()

def main():
    template_file_path = 'TMC.TMPL.docx'
    template_document = Document(template_file_path)

    handle_meeting_info(template_document)
    handle_agenda(template_document)

    output_file_path = 'TMC.{}.docx'.format(data_table['NO'])
    template_document.save(output_file_path)


if __name__ == '__main__':
    main()
